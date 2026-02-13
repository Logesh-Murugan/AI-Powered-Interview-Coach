"""
Text extraction utilities for PDF and DOCX files
"""
import re
import io
import requests
from typing import Optional, Tuple
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document
from loguru import logger


def clean_text(text: str) -> str:
    """
    Clean extracted text by removing extra whitespace and normalizing.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Replace multiple spaces with single space
    text = re.sub(r' +', ' ', text)
    
    # Replace multiple newlines with double newline
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove leading/trailing whitespace from entire text
    text = text.strip()
    
    return text


def extract_text_from_pdf_pypdf2(file_content: bytes) -> Optional[str]:
    """
    Extract text from PDF using PyPDF2.
    
    Args:
        file_content: PDF file content as bytes
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = '\n\n'.join(text_parts)
        
        # Check if we got meaningful text (not just whitespace)
        if full_text and len(full_text.strip()) > 50:
            return full_text
        
        return None
        
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {str(e)}")
        return None


def extract_text_from_pdf_pdfplumber(file_content: bytes) -> Optional[str]:
    """
    Extract text from PDF using pdfplumber (fallback method).
    
    Args:
        file_content: PDF file content as bytes
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        pdf_file = io.BytesIO(file_content)
        
        text_parts = []
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        full_text = '\n\n'.join(text_parts)
        
        # Check if we got meaningful text
        if full_text and len(full_text.strip()) > 50:
            return full_text
        
        return None
        
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {str(e)}")
        return None


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF with fallback strategy.
    
    Tries PyPDF2 first, falls back to pdfplumber if needed.
    
    Args:
        file_content: PDF file content as bytes
        
    Returns:
        Extracted text
        
    Raises:
        Exception: If both extraction methods fail
    """
    # Try PyPDF2 first (faster)
    logger.info("Attempting PDF extraction with PyPDF2")
    text = extract_text_from_pdf_pypdf2(file_content)
    
    if text:
        logger.info(f"PyPDF2 extraction successful: {len(text)} characters")
        return text
    
    # Fallback to pdfplumber (more robust)
    logger.info("PyPDF2 failed, trying pdfplumber")
    text = extract_text_from_pdf_pdfplumber(file_content)
    
    if text:
        logger.info(f"pdfplumber extraction successful: {len(text)} characters")
        return text
    
    # Both methods failed
    raise Exception("Failed to extract text from PDF using both PyPDF2 and pdfplumber")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_content: DOCX file content as bytes
        
    Returns:
        Extracted text
        
    Raises:
        Exception: If extraction fails
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = '\n\n'.join(text_parts)
        
        if not full_text or len(full_text.strip()) < 50:
            raise Exception("Extracted text is too short or empty")
        
        logger.info(f"DOCX extraction successful: {len(full_text)} characters")
        return full_text
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def download_file_from_url(url: str) -> bytes:
    """
    Download file from URL or read from local path.
    
    Args:
        url: File URL (can be HTTP URL or local path)
        
    Returns:
        File content as bytes
        
    Raises:
        Exception: If download/read fails
    """
    try:
        # Check if it's a local file path
        if url.startswith('/uploads/'):
            from pathlib import Path
            file_path = Path(url.lstrip('/'))
            
            if not file_path.exists():
                raise Exception(f"Local file not found: {file_path}")
            
            logger.info(f"Reading local file: {file_path}")
            with open(file_path, 'rb') as f:
                content = f.read()
            
            logger.info(f"Read {len(content)} bytes from local file")
            return content
        else:
            # HTTP URL - download from remote server
            logger.info(f"Downloading from URL: {url}")
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return response.content
            
    except Exception as e:
        logger.error(f"Failed to download/read file from {url}: {str(e)}")
        raise Exception(f"Failed to download file: {str(e)}")


def extract_text_from_resume(file_url: str, file_extension: str) -> Tuple[str, bool]:
    """
    Extract text from resume file (PDF or DOCX).
    
    Args:
        file_url: Local file URL or path
        file_extension: File extension (.pdf or .docx)
        
    Returns:
        Tuple of (extracted_text, success)
        
    Raises:
        Exception: If extraction fails
    """
    try:
        # Download file
        logger.info(f"Downloading file from {file_url}")
        file_content = download_file_from_url(file_url)
        logger.info(f"Downloaded {len(file_content)} bytes")
        
        # Extract text based on file type
        if file_extension.lower() == '.pdf':
            raw_text = extract_text_from_pdf(file_content)
        elif file_extension.lower() == '.docx':
            raw_text = extract_text_from_docx(file_content)
        else:
            raise Exception(f"Unsupported file extension: {file_extension}")
        
        # Clean text
        cleaned_text = clean_text(raw_text)
        
        if not cleaned_text or len(cleaned_text.strip()) < 50:
            raise Exception("Extracted text is too short after cleaning")
        
        logger.info(f"Text extraction successful: {len(cleaned_text)} characters")
        return cleaned_text, True
        
    except Exception as e:
        logger.error(f"Text extraction failed: {str(e)}")
        raise


def get_text_statistics(text: str) -> dict:
    """
    Get statistics about extracted text.
    
    Args:
        text: Extracted text
        
    Returns:
        Dictionary with statistics
    """
    if not text:
        return {
            'character_count': 0,
            'word_count': 0,
            'line_count': 0,
            'paragraph_count': 0
        }
    
    words = text.split()
    lines = text.split('\n')
    paragraphs = [p for p in text.split('\n\n') if p.strip()]
    
    return {
        'character_count': len(text),
        'word_count': len(words),
        'line_count': len(lines),
        'paragraph_count': len(paragraphs)
    }
